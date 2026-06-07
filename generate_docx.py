#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Стили страницы ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# Базовый шрифт
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Вспомогательные функции ─────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h4(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # серый фон через shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # заголовки
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shd)
    # строки
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

def pagebreak():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# ТИТУЛЬНАЯ СТРАНИЦА
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NeuroBalance Backend')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Полная техническая документация')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Версия 3.0  ·  Дата: 2026-06-07')
run.font.size = Pt(12)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Для диплома, разработки и защиты')
run.font.size = Pt(12)
run.font.italic = True

pagebreak()

# ════════════════════════════════════════════════════════════════
# СОДЕРЖАНИЕ
# ════════════════════════════════════════════════════════════════
h1('СОДЕРЖАНИЕ')

toc_items = [
    ('ЧАСТЬ I — АРХИТЕКТУРА СИСТЕМЫ', ''),
    ('  1.  Обзор и контекст проекта', ''),
    ('  2.  Микросервисная архитектура', ''),
    ('  3.  Технологический стек', ''),
    ('  4.  Как запустить проект', ''),
    ('', ''),
    ('ЧАСТЬ II — СЕРВИСЫ: API И БИЗНЕС-ЛОГИКА', ''),
    ('  5.  NBAuthService — Аутентификация (порт 8081)', ''),
    ('  6.  NBCheckinService — Чекины и здоровье (порт 8082)', ''),
    ('  7.  NoteAI-backend — Заметки с AI (порт 8083)', ''),
    ('', ''),
    ('ЧАСТЬ III — MACHINE LEARNING', ''),
    ('  8.  ML Service — Архитектура и роль', ''),
    ('  9.  Датасет и признаки (Feature Engineering)', ''),
    ('  10. Модели: обучение и гиперпараметры', ''),
    ('  11. Метрики качества и измерение ошибок', ''),
    ('  12. Формулы Health Metrics (M-Rest / M-Ready / M-Balance)', ''),
    ('  13. Почему XGBoost и CatBoost лучше альтернатив', ''),
    ('  14. Интеграция ML-сервиса в общую систему', ''),
    ('', ''),
    ('ЧАСТЬ IV — ИНФРАСТРУКТУРА', ''),
    ('  15. Kafka — Асинхронные события', ''),
    ('  16. Базы данных', ''),
    ('  17. JWT и безопасность', ''),
    ('', ''),
    ('ЧАСТЬ V — ИГРОВАЯ СИСТЕМА', ''),
    ('  18. XP, персонажи, стрики, бейджи', ''),
    ('', ''),
    ('ЧАСТЬ VI — ПОЛНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ', ''),
    ('  19. Примеры флоу от начала до конца', ''),
    ('', ''),
    ('ЧАСТЬ VII — ДЛЯ ЗАЩИТЫ ДИПЛОМА', ''),
    ('  20. Что рассказать на защите: бэкенд', ''),
    ('  21. Что рассказать на защите: ML часть', ''),
    ('  22. Сравнение с аналогами и научное обоснование', ''),
    ('  23. Структура слайдов для презентации', ''),
]
for item, _ in toc_items:
    if item == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith(' '):
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ I
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ I — АРХИТЕКТУРА СИСТЕМЫ')

# ── 1 ───────────────────────────────────────────────────────────
h2('1. Обзор и контекст проекта')

body(
    'NeuroBalance — это мобильное приложение для iOS, которое ежедневно отслеживает '
    'когнитивное благополучие пользователя. Каждое утро пользователь проходит короткий '
    'чекин (настроение, сон, стресс, физическая активность), после чего система '
    'автоматически выполняет четыре действия:'
)
bullet('Вычисляет три метрики здоровья: M-Rest, M-Ready и M-Balance.')
bullet('Предсказывает когнитивный скор от 0 до 100 через ML-модели (XGBoost).')
bullet('Генерирует персонализированные рекомендации на основе реальных данных пользователя (CatBoost).')
bullet('Обновляет геймификацию: XP персонажа, стрик и уровень.')

body(
    'Система закрывает задокументированный рыночный пробел: фитнес-трекеры измеряют '
    'физическую активность, но ни одно массовое приложение не предоставляет непрерывный, '
    'ML-управляемый индекс когнитивного благополучия из данных ежедневного самоотчёта. '
    'Глобальный рынок mental-wellness приложений оценивался в 5,2 млрд USD в 2023 году '
    'и растёт на 16,5% в год (Grand View Research, 2024).'
)
body('Целевая аудитория: работающие специалисты и студенты 22–40 лет в Казахстане и Центральной Азии.')

# ── 2 ───────────────────────────────────────────────────────────
h2('2. Микросервисная архитектура')

body(
    'Система построена по принципу микросервисной архитектуры и состоит из шести '
    'независимо развёртываемых компонентов, объединённых в один Docker Compose стек '
    'из десяти контейнеров. Каждый сервис имеет собственную изолированную базу данных '
    '(паттерн Database-per-Service, Newman 2015), что исключает schema coupling и '
    'позволяет каждому сервису эволюционировать независимо.'
)

add_table(
    ['Сервис', 'Порт', 'База данных', 'Роль'],
    [
        ['NBAuthService', '8081', 'auth_db (PostgreSQL)', 'Регистрация, JWT, онбординг, профиль'],
        ['NBCheckinService', '8082', 'checkin_db (PostgreSQL)', 'Чекины, сон, настроение, игры, XP, ML'],
        ['NoteAI-backend', '8083', 'noteai_db (PostgreSQL)', 'Заметки, дневник, AI-анализ через Groq'],
        ['ML Service', '5001', '— (in-memory cache)', 'XGBoost + CatBoost: когнитивный скор, рекомендации'],
        ['Apache Kafka', '29092', '—', 'Асинхронный обмен событиями'],
        ['Zookeeper', '2181', '—', 'Координация Kafka'],
    ]
)

body(
    'Перекрёстные ссылки между базами (например, userId в checkin_db) — это мягкие '
    'ссылки без FK-constraints, принимающие eventual consistency в обмен на '
    'независимую развёртываемость. ML Service доступен только внутри Docker-сети — '
    'никакой прямой доступ с фронтенда невозможен.'
)

body('Базовые URL для каждого сервиса:')
bullet('Auth Service:    http://localhost:8081/api/v1/')
bullet('CheckIn Service: http://localhost:8082/api/v1/')
bullet('NoteAI Service:  http://localhost:8083/api/v1/')
bullet('ML Service:      http://localhost:5001/')
bullet('Swagger UI доступен на каждом сервисе по пути /swagger-ui.html (ML: /apidocs)')

# ── 3 ───────────────────────────────────────────────────────────
h2('3. Технологический стек')

h3('3.1 Бэкенд (Java / Spring Boot)')

body(
    'Все три Java-сервиса написаны на Java 17 LTS (Long-Term Support) с использованием '
    'Spring Boot 4.0.2. Java 17 выбран за sealed classes, improved pattern matching и '
    'G1 Garbage Collector с минимальными паузами — что важно для heap-памяти '
    'типичного микросервисного контейнера (128–512 МБ). Spring Boot обеспечивает '
    'convention-over-configuration, встроенный Tomcat-сервер, dependency injection '
    'через application context и бесшовную интеграцию с security, JPA и Kafka.'
)

add_table(
    ['Компонент', 'Версия', 'Назначение'],
    [
        ['Spring Boot', '4.0.2', 'Основной фреймворк, embedded Tomcat, auto-configuration'],
        ['Java', '17 (LTS)', 'Runtime, sealed classes, G1 GC'],
        ['Spring Security 6', '6.x', 'Stateless JWT-аутентификация, BCrypt пароли'],
        ['Spring Data JPA / Hibernate', '6.x', 'ORM-слой, репозитории, JPQL-запросы'],
        ['HikariCP', '—', 'Пул соединений (max pool=10, timeout 30 сек)'],
        ['Apache Kafka', '7.6', 'Асинхронное расцепление сервисов'],
        ['springdoc-openapi', '2.8.8', 'Автогенерация Swagger UI на /swagger-ui.html'],
        ['JJWT', '0.12.6', 'Создание и валидация JWT (алгоритм HS256)'],
        ['Spring AI', '1.0.0-M6', 'Интеграция с Groq API для NoteAI (Llama 3.3-70b)'],
        ['Flyway', '—', 'Версионные миграции схемы БД (V1__, V2__...)'],
    ]
)

h3('3.2 ML Service (Python / Flask)')

body(
    'ML Service реализован на Python 3.11 с использованием Flask 3.0 в качестве '
    'лёгкого REST-фреймворка. Сервис недоступен снаружи Docker-сети — это намеренное '
    'архитектурное решение, исключающее прямой внешний доступ к ML-inference. '
    'Модели загружаются в память при старте Flask и хранятся там, обеспечивая '
    'нулевое дисковое I/O при каждом инференсе.'
)

add_table(
    ['Компонент', 'Версия', 'Назначение'],
    [
        ['Python', '3.11', 'ML runtime'],
        ['Flask', '3.0', 'REST API для ML inference'],
        ['XGBoost', '1.7.6', 'Регрессия когнитивного скора + классификация состояния'],
        ['CatBoost', '1.2', 'Ранжирование рекомендаций'],
        ['scikit-learn', '1.3.2', 'StandardScaler, LabelEncoder, метрики оценки'],
        ['Flasgger', '0.9.7', 'Swagger UI для ML-сервиса на /apidocs'],
        ['SHAP', '0.45', 'Интерпретируемость важности признаков'],
        ['joblib', '1.3', 'Сериализация моделей в .pkl файлы'],
        ['gunicorn', '21.2', 'Production WSGI-сервер для Flask'],
    ]
)

h3('3.3 iOS Frontend')
add_table(
    ['Компонент', 'Назначение'],
    [
        ['Swift 5.9 + SwiftUI', 'Декларативный UI, MVVM, iOS 16+'],
        ['URLSession async/await', 'Асинхронные сетевые запросы (Swift structured concurrency)'],
        ['Keychain API', 'Хранение JWT (шифрование через Secure Enclave устройства)'],
        ['Combine', 'Реактивная валидация форм с debounce'],
        ['Core Data + UserDefaults', 'Offline-first кэш для работы без сети'],
        ['UNUserNotificationCenter', 'Локальные уведомления (напоминание о чекине)'],
    ]
)

h3('3.4 Инфраструктура')
add_table(
    ['Компонент', 'Назначение'],
    [
        ['Docker Compose', '10 контейнеров: 3 Spring Boot + 3 PostgreSQL + ML + Kafka + Zookeeper + Swagger'],
        ['PostgreSQL 15', 'Три независимые БД (auth_db, checkin_db, noteai_db)'],
        ['DigitalOcean Droplet', 'Production-деплой всего стека'],
        ['pgAdmin', 'Веб-UI для администрирования PostgreSQL (порт 5050)'],
    ]
)

# ── 4 ───────────────────────────────────────────────────────────
h2('4. Как запустить проект')

body('Требования: Docker Desktop >= 4.x, Java 21 (для локальной разработки), Python 3.10+ (для ML без Docker).')

h3('4.1 Быстрый запуск через Docker Compose')
code(
    'cp .env.example .env\n'
    '# Заполнить JWT_SECRET, пароли БД, GROQ_API_KEY\n'
    'docker-compose up --build -d\n'
    'docker-compose ps'
)

h3('4.2 Порядок запуска контейнеров')
body('Docker Compose соблюдает зависимости автоматически:')
bullet('zookeeper → kafka')
bullet('postgres → nbauthservice')
bullet('postgres-checkin → nb-checkin-service')
bullet('ml-service → nb-checkin-service')
bullet('kafka → nb-checkin-service')
bullet('postgres-noteai → noteai-service')
bullet('nbauthservice → nb-checkin-service, noteai-service')

h3('4.3 Ключевые переменные окружения (.env)')
add_table(
    ['Переменная', 'Описание', 'Пример'],
    [
        ['JWT_SECRET', 'Секрет для подписи HS256 токенов', 'your-long-random-secret'],
        ['JWT_EXPIRATION', 'Срок действия accessToken (мс)', '86400000 (24 часа)'],
        ['AUTH_DB_PASSWORD', 'Пароль auth_db', 'secret'],
        ['CHECKIN_DB_PASSWORD', 'Пароль checkin_db', 'secret'],
        ['NOTEAI_DB_PASSWORD', 'Пароль noteai_db', 'secret'],
        ['GROQ_API_KEY', 'Ключ Groq API для Llama 3.3-70b', 'gsk_...'],
        ['MAIL_USERNAME / MAIL_PASSWORD', 'SMTP Gmail для email-верификации', 'user@gmail.com / app_password'],
        ['TWILIO_ACCOUNT_SID / AUTH_TOKEN', 'SMS-верификация через Twilio', '...'],
    ]
)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ II
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ II — СЕРВИСЫ: API И БИЗНЕС-ЛОГИКА')

# ── 5 ───────────────────────────────────────────────────────────
h2('5. NBAuthService — Аутентификация и пользователи (порт 8081)')

body('Base URL: http://localhost:8081/api/v1')

h3('5.1 Как работает авторизация')
body(
    'Все защищённые эндпойнты требуют JWT-токен в заголовке Authorization: Bearer <token>. '
    'Поток следующий: пользователь логинится и получает пару токенов — accessToken и '
    'refreshToken. При каждом запросе к NBCheckinService или NoteAI-backend '
    'JwtAuthenticationFilter (extends OncePerRequestFilter) валидирует токен, '
    'извлекает числовой userId из claim "id" и кладёт его в request.setAttribute("userId"). '
    'Контроллеры читают userId напрямую — SecurityContext не используется в CheckIn и NoteAI. '
    'Это stateless-архитектура.'
)
body('Параметры JWT:')
bullet('Алгоритм подписи: HS256 (HMAC-SHA-256), симметричный, секрет в JWT_SECRET')
bullet('Срок действия: 86 400 секунд (24 часа), настраивается через JWT_EXPIRATION')
bullet('Payload: id (userId), sub (username), roles (ROLE_USER/ROLE_ADMIN), exp')
bullet('Хранение паролей: BCrypt, cost factor 10 — соответствует NIST SP 800-63B')
bullet('На iOS: хранится в Keychain (kSecClassGenericPassword), шифрование Secure Enclave, OWASP MASVS-STORAGE-1')
bullet('Refresh-токен: UUID-строка в таблице users, используется для обновления accessToken без повторного логина')

h3('5.2 Основные эндпойнты регистрации и входа')

add_table(
    ['Метод', 'Путь', 'Авторизация', 'Описание'],
    [
        ['POST', '/auth/register', 'Нет', 'Регистрация. Онбординг можно передать сразу (опционально)'],
        ['POST', '/auth/login', 'Нет', 'Вход, возвращает accessToken + refreshToken'],
        ['POST', '/auth/refresh', 'Нет', 'Обновить accessToken, тело — raw refresh token'],
        ['GET', '/auth/verify-email?token=uuid', 'Нет', 'Подтверждение email по ссылке из письма'],
        ['POST', '/auth/verify-user-manual?username=x', 'Нет', 'Ручная верификация (для тестов)'],
        ['POST', '/auth/register-phone', 'Нет', 'Регистрация по телефону — SMS через Twilio'],
        ['POST', '/auth/verify-phone?phone=+7...&code=123456', 'Нет', 'Верификация SMS кода'],
    ]
)

body('Пример тела запроса для POST /auth/register:')
code(
    '{\n'
    '  "name": "Аман Гельди",\n'
    '  "email": "user@example.com",\n'
    '  "username": "amangeldi",\n'
    '  "phone": "+77001234567",\n'
    '  "password": "Password123",\n'
    '  "sex": "male",\n'
    '  "heightCm": 175,\n'
    '  "weightKg": 70.5,\n'
    '  "birthDate": "2001-04-15",\n'
    '  "characterId": 1,\n'
    '  "dataConsent": true\n'
    '}'
)
body('Ответ при успешной регистрации с онбордингом (200 OK):')
code(
    '{\n'
    '  "message": "Registration successful",\n'
    '  "userId": 42,\n'
    '  "username": "amangeldi",\n'
    '  "isOnboarded": true,\n'
    '  "onboardingCompleted": true\n'
    '}'
)

h3('5.3 Онбординг (/onboarding) — все эндпойнты требуют JWT')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['POST', '/onboarding', 'Создать или полностью обновить данные онбординга'],
        ['GET', '/onboarding', 'Получить данные онбординга текущего пользователя'],
        ['GET', '/onboarding/status', 'Проверить: isOnboarded и isCompleted'],
        ['PATCH', '/onboarding', 'Частично обновить (только переданные поля)'],
        ['DELETE', '/onboarding', 'Удалить данные онбординга'],
        ['GET', '/onboarding/user/{userId}', 'Получить онбординг конкретного юзера (admin)'],
    ]
)

h3('5.4 Профиль пользователя (/users)')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['GET', '/users/me', 'Полный профиль: id, name, email, username, phone, isOnboarded, onboarding'],
        ['GET', '/users/{id}', 'Пользователь по ID'],
        ['PUT', '/users', 'Обновить профиль (обязательно передать id)'],
        ['GET', '/users/search?q=query', 'Поиск по имени/username — возвращает id, username, name'],
        ['DELETE', '/users/{id}', 'Удалить пользователя'],
    ]
)

# ── 6 ───────────────────────────────────────────────────────────
h2('6. NBCheckinService — Чекины и данные здоровья (порт 8082)')

body('Base URL: http://localhost:8082/api/v1')
body(
    'Все эндпойнты (кроме /checkins/health) требуют заголовок Authorization: Bearer <token>. '
    'Временная зона для всех дат "сегодня": Asia/Almaty (UTC+5).'
)

h3('6.1 Ежедневный чекин (/checkins)')
body(
    'Чекин — ежедневная форма данных о самочувствии пользователя. '
    'Разрешён один чекин в день. После создания через Kafka автоматически '
    'запускается пересчёт Health Metrics и ML-рекомендаций. '
    'Задача COMPLETE_CHECKIN отмечается выполненной (+50 XP персонажу). Стрик обновляется.'
)

body('Обязательные поля при создании чекина (POST /checkins):')
add_table(
    ['Поле', 'Тип', 'Диапазон', 'Обязательно'],
    [
        ['sleepQuality', 'int', '1–10', 'Да'],
        ['energyLevel', 'int', '1–10', 'Да'],
        ['stressLevel', 'int', '1–10', 'Да'],
        ['morningMood', 'int', '1–5', 'Нет'],
        ['eveningMood', 'int', '1–5', 'Нет'],
        ['sleepHours', 'decimal', '0–24', 'Нет'],
        ['sleepBedtime / sleepWaketime', 'time', 'HH:mm', 'Нет'],
        ['physicalActivityMinutes', 'int', '0+', 'Нет'],
        ['physicalActivityType', 'string', 'walk/gym/yoga/sports/none', 'Нет'],
        ['didExercise, ateHealthy, hadSocialInteraction', 'boolean', '—', 'Нет'],
    ]
)

add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['POST', '/checkins', 'Создать чекин на сегодня'],
        ['GET', '/checkins/today', 'Сегодняшний чекин (404 если не создан)'],
        ['GET', '/checkins/today/exists', 'Был ли чекин сегодня: { exists, canCheckIn }'],
        ['GET', '/checkins/{date}', 'Чекин за дату (yyyy-MM-dd)'],
        ['PUT', '/checkins/{date}', 'Обновить чекин за дату'],
        ['DELETE', '/checkins/{date}', 'Удалить чекин за дату'],
        ['GET', '/checkins/recent', 'Последние 30 дней'],
        ['GET', '/checkins?startDate=...&endDate=...', 'Диапазон дат'],
        ['GET', '/checkins/calendar?year=2026&month=6', 'Даты выполненных чекинов (для календаря)'],
        ['GET', '/checkins/stats/weekly', 'Статистика за неделю: кол-во, средние настроение/сон/стресс'],
        ['GET', '/checkins/stats/monthly', 'Статистика за месяц'],
        ['GET', '/checkins/stats?startDate=...&endDate=...', 'Произвольный период'],
        ['GET', '/checkins/streak', 'Информация о стрике'],
        ['POST', '/checkins/streak/recalculate', 'Пересчитать стрик вручную'],
    ]
)

h3('6.2 Настроение (/mood) — несколько записей в день')
body('POST /mood автоматически отмечает задачу LOG_MOOD выполненной (+20 XP).')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['POST', '/mood', 'Записать настроение: moodScore(1-5), note, triggers[], loggedAt'],
        ['GET', '/mood', 'Все логи настроения'],
        ['GET', '/mood/recent', 'Последние 20 логов'],
        ['GET', '/mood/today', 'Логи за сегодня'],
        ['GET', '/mood/date/{date}', 'По конкретной дате (yyyy-MM-dd)'],
        ['GET', '/mood/range?startTime=...&endTime=...', 'Диапазон дат/времени'],
        ['GET', '/mood/{id}', 'По ID'],
        ['PUT', '/mood/{id}', 'Обновить лог'],
        ['DELETE', '/mood/{id}', 'Удалить лог'],
        ['GET', '/mood/average?startTime=...&endTime=...', 'Среднее настроение за период'],
        ['GET', '/mood/trigger/{trigger}', 'Логи по конкретному триггеру (work, sleep и т.д.)'],
    ]
)

h3('6.3 Сон (/sleep)')
body('POST /sleep автоматически отмечает LOG_SLEEP выполненной (+25 XP) и публикует Kafka-событие sleep.logged.')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['POST', '/sleep', 'Записать сон: sleepDate, bedtime, wakeTime, durationHours, qualityScore, deepSleepMinutes, remSleepMinutes'],
        ['GET', '/sleep', 'Все записи'],
        ['GET', '/sleep/recent', 'Последние записи'],
        ['GET', '/sleep/today', 'Запись за сегодня (404 если нет)'],
        ['GET', '/sleep/today/exists', 'Была ли запись: { exists, canLog }'],
        ['GET', '/sleep/{id}', 'По ID'],
        ['GET', '/sleep/date/{date}', 'По дате'],
        ['GET', '/sleep/range?startDate=...&endDate=...', 'Диапазон'],
        ['PUT', '/sleep/{id}', 'Обновить по ID'],
        ['PUT', '/sleep/date/{date}', 'Обновить по дате'],
        ['DELETE', '/sleep/{id}', 'Удалить по ID'],
        ['DELETE', '/sleep/date/{date}', 'Удалить по дате'],
    ]
)

h3('6.4 Ежедневные задачи (/tasks)')
body(
    'Каждый день для пользователя автоматически создаётся набор из пяти задач. '
    'Каждая задача выполняется автоматически при обращении к соответствующему эндпойнту — '
    'вручную отмечать не требуется.'
)
add_table(
    ['TaskType', 'Что нужно сделать', 'Автовыполнение при', 'XP'],
    [
        ['COMPLETE_CHECKIN', 'Ежедневный чекин', 'POST /checkins', '+50'],
        ['LOG_MOOD', 'Записать настроение', 'POST /mood', '+20'],
        ['LOG_SLEEP', 'Записать сон', 'POST /sleep', '+25'],
        ['PLAY_GAME', 'Сыграть в когнитивную игру', 'POST /brain-games/submit', '+30–80'],
        ['WRITE_NOTE', 'Написать заметку/дневник', 'POST /api/v1/journal', '+45'],
    ]
)
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['GET', '/tasks/today', 'Задачи на сегодня с isCompleted, completedAt, xpReward'],
        ['POST', '/tasks/complete?taskType=LOG_MOOD&date=...', 'Ручное выполнение задачи'],
        ['GET', '/tasks/stats', 'Статистика выполнения задач'],
        ['GET', '/tasks/date/{date}', 'Задачи за дату (создаёт если нет)'],
        ['GET', '/tasks/history?startDate=...&endDate=...', 'История задач'],
        ['POST', '/tasks/note-written?date=...', 'Отметить WRITE_NOTE — вызывается автоматически из NoteAI'],
    ]
)

h3('6.5 Метрики здоровья (/health-metrics)')
body(
    'Метрики вычисляются автоматически через ML Service после каждого чекина. '
    'Три метрики (M-Rest, M-Ready, M-Balance) и итоговый overall score. '
    'Метки: Excellent (>=80), Good (>=60), Fair (>=40), Poor (<40). '
    'Формулы подробно описаны в разделе 12.'
)
body('Пример ответа GET /health-metrics/today:')
code(
    '{\n'
    '  "userId": 42, "date": "2026-06-06",\n'
    '  "mRest": 78.5,    "mRestLabel": "Good",\n'
    '  "mReady": 72.3,   "mReadyLabel": "Good",\n'
    '  "mBalance": 65.1, "mBalanceLabel": "Good",\n'
    '  "overall": 71.9,  "overallLabel": "Good"\n'
    '}'
)
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['GET', '/health-metrics/today', 'Метрики за сегодня'],
        ['GET', '/health-metrics/{date}', 'Метрики за конкретную дату'],
        ['GET', '/health-metrics/recent?days=7', 'Последние N дней (1–90)'],
        ['GET', '/health-metrics/history', 'Полная история'],
        ['POST', '/health-metrics/recalculate?date=...', 'Принудительный пересчёт через ML'],
    ]
)

h3('6.6 ML-рекомендации (/ml)')
body(
    'Персонализированные рекомендации обновляются автоматически через Kafka '
    'при каждом чекине, сне или завершённой игре. '
    'Все основные параметры (сон, стресс, возраст, пол) берутся из реальных данных '
    'пользователя в БД — тело запроса служит только для дополнительных предпочтений.'
)
body('Пример ответа POST /ml/recommendations:')
code(
    '{\n'
    '  "userId": 42,\n'
    '  "cognitiveScore": 68.3,\n'
    '  "cognitiveState": "Medium",\n'
    '  "recommendations": [\n'
    '    {\n'
    '      "type": "SLEEP_INCREASE",\n'
    '      "title": "Оптимизация сна",\n'
    '      "priority": "HIGH",\n'
    '      "predictedImprovement": 8.5,\n'
    '      "actions": ["Будильник на отбой в 23:00", "Затемните спальню"],\n'
    '      "scientificBasis": "Walker (2017): Why We Sleep",\n'
    '      "baseline": 6.0,\n'
    '      "recommendedTarget": 7.5\n'
    '    }\n'
    '  ],\n'
    '  "totalPotential": 22.4,\n'
    '  "summary": "Ваш когнитивный счет: 68.3. Рекомендации могут дать +22.4 балла."\n'
    '}'
)
body('Четыре типа рекомендаций: SLEEP_INCREASE, STRESS_DECREASE, EXERCISE_INCREASE, SCREEN_DECREASE.')

h3('6.7 Стрики (/streaks), персонаж (/character), игры (/brain-games, /game-sessions)')
add_table(
    ['Эндпойнт', 'Описание'],
    [
        ['GET /streaks', 'currentStreak, longestStreak, totalCheckins, totalXpEarned, nextMilestone'],
        ['GET /streaks/leaderboard/streak', 'Топ-10 пользователей по стрику'],
        ['GET /streaks/leaderboard/xp', 'Топ-10 по XP'],
        ['GET /character', 'Текущий персонаж: type, level, xp, xpToNextLevel, happiness, energy'],
        ['POST /character/select', 'Выбрать персонажа (CAT, DOG, RABBIT, PANDA, FOX)'],
        ['POST /character/feed', 'Покормить (+10 счастья, +15 энергии)'],
        ['POST /character/play', 'Поиграть (+15 счастья, -5 энергии)'],
        ['POST /brain-games/submit', 'Отправить результат: gameType, score, durationSeconds, difficultyLevel, correctAnswers'],
        ['GET /brain-games/stats', 'Агрегированная статистика игр'],
        ['POST /brain-games/level-up', 'Повысить уровень персонажа (по недельным результатам)'],
        ['POST /game-sessions', 'Fun-игры (DONUT_GAME, CHARACTER_CARE): score, durationSeconds, isCompleted'],
        ['POST /new-game-sessions', 'DONUT_GAME, NUMBER_SEQUENCE_GAME с attemptsCount'],
        ['GET /rewards', 'Все бейджи: isUnlocked, progress, target, xpBonus, xpMultiplier'],
        ['POST /rewards/check', 'Проверить и разблокировать новые награды'],
    ]
)

# ── 7 ───────────────────────────────────────────────────────────
h2('7. NoteAI-backend — Заметки с AI (порт 8083)')

body('Base URL: http://localhost:8083')
body(
    'AI-анализ работает через Groq Cloud API с моделью llama-3.3-70b-versatile. '
    'Параметр temperature = 0.3 — для детерминированных, фактических ответов. '
    'Groq использует LPU (Language Processing Unit) — специализированное железо '
    'для LLM-инференса, обеспечивающее latency менее 1 секунды. '
    'Интеграция реализована через Spring AI (spring-ai-openai-spring-boot-starter v1.0.0-M6).'
)

h3('7.1 Дневник (/api/v1/journal)')
body(
    'Основной модуль для ведения дневника. Поддерживает auto-save через PATCH с дебаунсом 2–3 секунды. '
    'При создании записи автоматически выполняется задача WRITE_NOTE (+45 XP) через NBCheckinService.'
)
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['POST', '/api/v1/journal', 'Создать запись: title, content, moodScore, tags, isFavorite'],
        ['GET', '/api/v1/journal', 'Все записи (от новых к старым)'],
        ['GET', '/api/v1/journal/today', 'Записи за сегодня (Asia/Almaty)'],
        ['GET', '/api/v1/journal/date/{date}', 'За конкретную дату'],
        ['GET', '/api/v1/journal/range?startDate=...&endDate=...', 'Диапазон дат'],
        ['GET', '/api/v1/journal/favorites', 'Только isFavorite=true'],
        ['GET', '/api/v1/journal/{id}', 'По ID'],
        ['PUT', '/api/v1/journal/{id}', 'Полное обновление (WRITE_NOTE не дублируется)'],
        ['PATCH', '/api/v1/journal/{id}', 'Частичное обновление / auto-save (только изменённые поля)'],
        ['DELETE', '/api/v1/journal/{id}', 'Удалить запись'],
        ['POST', '/api/v1/journal/{id}/ai/summary', 'AI резюме: 1–2 предложения на языке записи'],
        ['POST', '/api/v1/journal/{id}/ai/analyze', 'AI wellness-анализ: tone, themes, wellnessInsight, suggestion'],
        ['POST', '/api/v1/journal/ai/chat', 'AI чат-ассистент с опциональным noteId для контекста'],
    ]
)
body('Пример ответа AI-анализа (POST /api/v1/journal/{id}/ai/analyze):')
code(
    '{\n'
    '  "noteId": 15, "type": "analysis",\n'
    '  "tone": "тревожный",\n'
    '  "themes": ["рабочий стресс", "усталость", "нехватка сна"],\n'
    '  "wellnessInsight": "Высокий стресс снижает M-Balance ниже 50 и давит на M-Ready.",\n'
    '  "suggestion": "Прогулка 10 минут после обеда снижает кортизол на 15-20%.",\n'
    '  "generatedAt": "2026-06-06 14:35:00"\n'
    '}'
)

h3('7.2 Заметки — Notes (/api/v1/notes)')
body('Отдельный модуль заметок с отдельной таблицей от Journal. Те же AI-методы (summary, analyze, chat).')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['GET', '/api/v1/notes', 'Все заметки'],
        ['GET', '/api/v1/notes/{id}', 'По ID'],
        ['PUT', '/api/v1/notes/{id}', 'Обновить (title + content)'],
        ['DELETE', '/api/v1/notes/{id}', 'Удалить'],
        ['POST', '/api/v1/notes/{id}/ai/summary', 'AI резюме'],
        ['POST', '/api/v1/notes/{id}/ai/analyze', 'AI анализ'],
        ['POST', '/api/v1/notes/ai/chat', 'AI чат'],
    ]
)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ III — ML
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ III — MACHINE LEARNING')

# ── 8 ───────────────────────────────────────────────────────────
h2('8. ML Service — Архитектура и роль')

body(
    'ML Service — отдельный микросервис на Python/Flask (версия 5.1.0 "Ultimate Edition"), '
    'решающий три принципиально разные задачи: регрессию (предсказание когнитивного скора), '
    'классификацию (определение когнитивного состояния) и ранжирование (выбор лучших '
    'персонализированных рекомендаций). Сервис недоступен снаружи Docker-сети — '
    'исключительно внутренний вызов из Java-бэкенда. Это архитектурное решение '
    'предотвращает прямой внешний доступ к ML-инференсу.'
)

body('Шесть моделей в системе:')
add_table(
    ['Задача', 'Алгоритм', 'Эндпойнт', 'Файл модели'],
    [
        ['Когнитивный скор (0–100)', 'XGBoost Regressor', 'POST /predict', 'cognitive_score_model.pkl'],
        ['Когнитивное состояние (HIGH/MEDIUM/LOW)', 'XGBoost Classifier', 'POST /predict', 'cognitive_state_model.pkl'],
        ['Рекомендации (выбор и ранжирование)', 'CatBoost Regressor', 'POST /recommend/top3', 'catboost_recommendation_model.cbm'],
        ['M-Rest (0–100)', 'XGBoost Regressor', 'POST /health-metrics/calculate', 'hm_m_rest_model.pkl'],
        ['M-Ready (0–100)', 'XGBoost Regressor', 'POST /health-metrics/calculate', 'hm_m_ready_model.pkl'],
        ['M-Balance (0–100)', 'XGBoost Regressor', 'POST /health-metrics/calculate', 'hm_m_balance_model.pkl'],
    ]
)

body(
    'Все файлы моделей загружаются в память при старте Flask — нулевое дисковое I/O при инференсе. '
    'Также в сервисе есть in-memory per-user кэш (_user_cache), ключ — str(user_id), '
    'значение — результат рекомендаций с датой (Asia/Almaty). При наступлении нового дня '
    'кэш автоматически инвалидируется и пересчитывается.'
)

# ── 9 ───────────────────────────────────────────────────────────
h2('9. Датасет и признаки (Feature Engineering)')

h3('9.1 Исходный датасет для когнитивного скора')
body(
    'Файл: ml/human_cognitive_performance.csv. '
    'Датасет содержит физиологически валидированные данные о когнитивной '
    'производительности людей с описанием их образа жизни. '
    'Используется для обучения XGBoost-моделей.'
)
add_table(
    ['Признак', 'Диапазон', 'Описание'],
    [
        ['Sleep_Duration', '4–10 ч', 'Продолжительность сна'],
        ['Stress_Level', '1–10', 'Уровень стресса'],
        ['Daily_Screen_Time', '1–14 ч', 'Экранное время'],
        ['Exercise_Frequency_Num', '0–7', 'Тренировок в неделю'],
        ['Caffeine_Intake', '0–6 порций', 'Потребление кофеина'],
        ['Reaction_Time', '200–500 мс', 'Время реакции (когнитивный тест)'],
        ['Memory_Test_Score', '40–100', 'Результат теста памяти'],
        ['Age', '18–65', 'Возраст'],
        ['Gender_Encoded', '0/1', 'Пол (Label Encoded)'],
        ['Diet_Non-Vegetarian / Diet_Vegan / Diet_Vegetarian', '0/1', 'One-hot кодировка типа питания'],
    ]
)

h3('9.2 Feature Engineering — 7 производных признаков')
body(
    'Feature engineering — ключевой шаг. Производные признаки улавливают взаимодействия, '
    'которые сырые признаки не могут выразить индивидуально. Итого: 17 финальных признаков (FINAL_FEATURES).'
)

h4('CFI — Cognitive Fatigue Index (индекс когнитивного утомления)')
body('Взвешенная комбинация четырёх факторов усталости:')
code(
    'r_norm     = (reaction_time - 200) / 400        # нормализованная реакция\n'
    's_norm     = (stress_level - 1) / 9              # нормализованный стресс\n'
    'sleep_debt = max(0, 7 - sleep_duration) / 3      # недосып от нормы\n'
    'screen_fat = max(0, daily_screen_time - 8) / 4   # избыток экранного времени\n'
    '\n'
    'CFI = (0.30 * r_norm + 0.25 * s_norm\n'
    '      + 0.25 * sleep_debt + 0.20 * screen_fat) * 100'
)
body('Основание: исследования Killgore (2010) по когнитивному утомлению. Веса: 30% реакция, 25% стресс, 25% недосып, 20% экраны.')

h4('Sleep_Debt — дефицит сна')
code('sleep_debt = max(0, 7 - sleep_duration)   # часы недосыпа относительно нормы 7 ч')

h4('Memory_Efficiency — эффективность памяти с учётом реакции')
code('memory_efficiency = (memory_test_score / reaction_time) * 1000')

h4('Lifestyle_Balance — общий баланс образа жизни (0–100)')
code(
    'sleep_sc = clip((sleep_duration - 4) / 6, 0, 1)\n'
    'ex_sc    = clip(exercise_frequency / 7, 0, 1)\n'
    'st_sc    = 1 - ((stress_level - 1) / 9)\n'
    'scr_sc   = 1 - clip((daily_screen_time - 1) / 11, 0, 1)\n'
    '\n'
    'lifestyle_balance = (0.30*sleep_sc + 0.25*st_sc\n'
    '                   + 0.20*ex_sc  + 0.15*scr_sc) * 100'
)

h4('Sleep_Exercise_Interaction, Sleep_Performance_Ratio, Stress_Resilience')
code(
    'sleep_exercise_interaction = sleep_duration * exercise_frequency\n'
    'sleep_performance_ratio   = memory_test_score / max(sleep_duration, 1)\n'
    'stress_resilience          = memory_test_score * (1 - stress_level / 10)'
)

h3('9.3 Датасет для обучения рекомендаций')
body(
    'Файл: ml/notebooks/recommendation_training_data_v3.csv. '
    'Около 50 000 записей. Для каждого профиля пользователя и каждого типа рекомендации '
    'рассчитано поле improvement — насколько вырастет когнитивный скор, '
    'если пользователь выполнит рекомендацию.'
)
add_table(
    ['Столбец', 'Описание'],
    [
        ['original_score', 'Базовый когнитивный скор пользователя'],
        ['improvement', 'Предсказанный прирост скора при выполнении рекомендации'],
        ['action_type', 'Тип: SLEEP_INCREASE / STRESS_DECREASE / EXERCISE_INCREASE / SCREEN_DECREASE'],
        ['action_magnitude', 'Величина изменения (например, +1.5 часа сна)'],
        ['baseline_value', 'Текущее значение параметра у пользователя'],
        ['is_critical', '1 если параметр в критической зоне (недосып <6ч, стресс >7 и т.д.)'],
        ['is_optimal_zone', '1 если параметр уже в оптимальном диапазоне'],
    ]
)

h3('9.4 Датасет для Health Metrics')
body(
    'Файл: ml/notebooks/cognitive_performance_engineered.csv. '
    '12 000 синтетических записей. Target values (M-Rest, M-Ready, M-Balance) '
    'рассчитаны по физиологически валидированным формулам (Stults-Kolehmainen & Sinha, 2014). '
    'Признаки: sleep_hours, sleep_quality, energy_level, morning_mood, evening_mood, '
    'stress_level, did_exercise, ate_healthy, had_social_interaction, felt_rested, '
    'deep_sleep_minutes, rem_sleep_minutes, total_sleep_minutes + 4 производных '
    '(sleep_debt, sleep_efficiency, wellbeing_index, habit_count).'
)

# ── 10 ──────────────────────────────────────────────────────────
h2('10. Модели: обучение и гиперпараметры')

h3('10.1 XGBoost Regressor — когнитивный скор')
body(
    'XGBoost (eXtreme Gradient Boosting) строит ансамбль неглубоких деревьев решений '
    'аддитивно: каждое новое дерево обучается исправлять остаточные ошибки предыдущего ансамбля. '
    'Функция потерь для регрессии — MSE с L1 (reg_alpha) и L2 (reg_lambda) регуляризацией '
    'для предотвращения переобучения. Ссылка: Chen & Guestrin (2016).'
)
code(
    'XGBRegressor(\n'
    '    n_estimators=500,        # количество деревьев\n'
    '    max_depth=8,             # глубина дерева\n'
    '    learning_rate=0.05,      # шаг обучения (eta)\n'
    '    subsample=0.8,           # доля наблюдений на дерево (защита от overfitting)\n'
    '    colsample_bytree=0.8,    # доля признаков на дерево\n'
    '    gamma=0.1,               # минимальный gain для сплита\n'
    '    reg_alpha=0.1,           # L1-регуляризация\n'
    '    reg_lambda=1.0,          # L2-регуляризация\n'
    '    random_state=42,\n'
    '    n_jobs=-1,               # все CPU-ядра\n'
    '    early_stopping_rounds=30 # стоп при отсутствии улучшения val RMSE\n'
    ')'
)
body('Разбивка данных: 70% train / 10% validation / 20% test, стратифицированная по Cognitive_State.')
body('Нормализация: StandardScaler — fit ТОЛЬКО на train, transform на val/test. Это защита от data leakage.')

h3('10.2 XGBoost Classifier — когнитивное состояние')
body(
    'Трёхклассовая классификация: HIGH, MEDIUM, LOW. '
    'Метки кодируются LabelEncoder (HIGH=0, LOW=1, MEDIUM=2), '
    'после предсказания — обратная трансформация inverse_transform.'
)
code(
    'XGBClassifier(\n'
    '    n_estimators=200,\n'
    '    max_depth=6,\n'
    '    learning_rate=0.1,\n'
    '    subsample=0.8,\n'
    '    colsample_bytree=0.8,\n'
    '    random_state=42,\n'
    '    early_stopping_rounds=20\n'
    ')'
)

h3('10.3 CatBoost Regressor — рекомендации')
body(
    'CatBoost (Categorical Boosting) — градиентный бустинг с нативной поддержкой '
    'категориальных признаков через Ordered Target Statistics (Prokhorenkova et al., 2018). '
    'Это исключает target leakage при обработке поля action_type. '
    'Использует use_best_model=True — автовыбор лучшей итерации по validation RMSE.'
)
code(
    'CatBoostRegressor(\n'
    '    iterations=1000,\n'
    '    learning_rate=0.05,\n'
    '    depth=8,\n'
    '    loss_function="RMSE",\n'
    '    eval_metric="RMSE",\n'
    '    random_seed=42,\n'
    '    use_best_model=True     # лучшая итерация по val set\n'
    ')'
)
body(
    'Входные признаки CatBoost: 17 признаков пользователя + '
    'one-hot типа рекомендации (4 столбца: action_SLEEP_INCREASE и т.д.) + '
    'action_delta, baseline_value, baseline_score, is_critical, is_optimal_zone = итого 24 признака.'
)
body(
    'SmartRecommendationSelector: для каждого из 4 типов рекомендаций и нескольких '
    'сценариев action_magnitude CatBoost предсказывает improvement. '
    'Выбирается лучший сценарий, дедупликация по типу, '
    'сортировка по predicted_improvement (descending), '
    'минимальный порог 0.45 для подавления слабых рекомендаций.'
)

h3('10.4 XGBoost × 3 — Health Metrics (M-Rest, M-Ready, M-Balance)')
body(
    'Три отдельных XGBoost Regressor — по одной модели для каждой метрики. '
    'Обучены на 12 000 синтетических записей. '
    'У каждой модели свой StandardScaler (hm_m_rest_scaler.pkl и т.д.). '
    'Graceful degradation: если модели недоступны — вычисляется по детерминированным '
    'формулам (см. раздел 12) с флагом source: "formula_fallback" в ответе.'
)

# ── 11 ──────────────────────────────────────────────────────────
h2('11. Метрики качества и измерение ошибок')

h3('11.1 Используемые метрики')
add_table(
    ['Метрика', 'Формула', 'Что показывает', 'Хорошее значение'],
    [
        ['R² (коэффициент детерминации)', '1 - SS_res / SS_tot', 'Доля объяснённой дисперсии. 1 = идеал, 0 = модель не лучше среднего', '> 0.90'],
        ['RMSE', 'sqrt(sum((y-ŷ)²)/n)', 'Среднеквадратичное отклонение в единицах таргета', 'Как можно меньше'],
        ['MAE', 'sum(|y-ŷ|)/n', 'Среднее абсолютное отклонение (менее чувствителен к выбросам)', 'Как можно меньше'],
        ['Accuracy', '(TP+TN)/Total', 'Доля верных предсказаний классификатора', '> 0.90'],
    ]
)

h3('11.2 Достигнутые результаты')
add_table(
    ['Модель', 'Метрика 1', 'Метрика 2', 'Интерпретация'],
    [
        ['XGBoost Regressor (когнитивный скор)', 'R² = 0.9990', 'RMSE = 0.73', 'Ошибка < 1 балла из 100'],
        ['XGBoost Classifier (состояние)', 'Accuracy = 97.9%', '—', '97.9% верных классификаций HIGH/MEDIUM/LOW'],
        ['CatBoost (рекомендации)', 'R² = 0.9656', 'RMSE = 1.81', 'Ошибка предсказания improvement < 2 баллов'],
        ['XGBoost M-Rest', 'R² = 0.917', 'MAE = 2.01', 'Средняя ошибка 2 балла из 100'],
        ['XGBoost M-Ready', 'R² = 0.971', 'MAE = 2.02', 'Средняя ошибка 2 балла из 100'],
        ['XGBoost M-Balance', 'R² = 0.967', 'MAE = 2.13', 'Средняя ошибка 2 балла из 100'],
    ]
)

h3('11.3 Защита от переобучения')
bullet(
    'Early stopping: обучение останавливается автоматически если validation RMSE '
    'не улучшается 30 итераций (30 для регрессора, 20 для классификатора).'
)
bullet(
    'Стратифицированный split 70/10/20: test set полностью изолирован от '
    'выбора гиперпараметров и early stopping.'
)
bullet(
    'StandardScaler: fit только на train — transform на val/test. '
    'Это стандартная защита от утечки информации из тестовой выборки.'
)
bullet(
    'L1 (reg_alpha=0.1) и L2 (reg_lambda=1.0) регуляризация в XGBoost: '
    'штрафуют сложность деревьев, предотвращая overfit на синтетических данных.'
)

h3('11.4 Корректность на продакшне')
bullet('Graceful degradation: при недоступности ML Java вычисляет метрики по формулам.')
bullet('Inference latency: XGBoost — 2–5 мс, CatBoost — 10–20 мс (всё в памяти, без disk I/O).')
bullet('Кэш инвалидируется в 00:00 Asia/Almaty через Kafka-события — данные всегда актуальны.')
bullet('SHAP (SHapley Additive exPlanations) используется для интерпретации важности признаков.')

# ── 12 ──────────────────────────────────────────────────────────
h2('12. Формулы Health Metrics (M-Rest / M-Ready / M-Balance)')

body(
    'Метрики имеют два режима вычисления: детерминированные формулы (быстрые, всегда работают) '
    'и ML-предсказание через XGBoost (точнее, если модели загружены). '
    'При доступности ML-сервиса используется значение из ML.'
)

h3('12.1 M-Rest — качество восстановления и сна (0–100)')
code(
    'M-Rest = min(sleep_hours / 8, 1.0) * 50\n'
    '       + (sleep_quality / 10.0)    * 30\n'
    '       + felt_rested               * 20\n'
    '\n'
    '# Зажато в [0, 100]'
)
body('Альтернативная формула (детальная, из ML-сервиса):')
code(
    'M-Rest = (sleepHours / 8.0)              * 40\n'
    '       + (sleepQuality / 10.0)           * 40\n'
    '       + (deepSleepMinutes / 90.0)       * 20'
)
add_table(
    ['Компонент', 'Вес', 'Научное основание'],
    [
        ['sleep_hours / 8 (норма)', '50%', 'Walker (2017): норма сна 7–9 часов'],
        ['sleep_quality / 10', '30%', 'Killgore (2010): субъективное качество коррелирует с deep sleep'],
        ['felt_rested (субъективно)', '20%', 'Субъективное ощущение восстановления'],
    ]
)

h3('12.2 M-Ready — когнитивная готовность к дню (0–100)')
code(
    'M-Ready = (energy_level / 10.0)   * 40\n'
    '        + (morning_mood / 5.0)    * 30\n'
    '        + M-Rest                  * 0.30'
)
add_table(
    ['Компонент', 'Вес', 'Обоснование'],
    [
        ['energy_level', '40%', 'Главный предиктор готовности к когнитивным задачам'],
        ['morning_mood', '30%', 'Утреннее настроение = прокси когнитивного состояния'],
        ['M-Rest', '30%', 'Качество вчерашнего сна определяет сегодняшнюю продуктивность'],
    ]
)

h3('12.3 M-Balance — эмоциональный и поведенческий баланс (0–100)')
code(
    'M-Balance = max((10 - stress_level) / 9, 0)  * 40\n'
    '          + did_exercise               * 13.3\n'
    '          + ate_healthy                * 13.3\n'
    '          + had_social_interaction     * 13.4\n'
    '          + (evening_mood / 5.0)       * 20'
)
add_table(
    ['Компонент', 'Вес', 'Обоснование'],
    [
        ['stress_level (инверсный)', '40%', 'McEwen (2007): стресс = основной разрушитель эмоционального баланса'],
        ['lifestyle habits (3 штуки)', '40%', 'Stults-Kolehmainen & Sinha (2014): упражнения + питание + социальность'],
        ['evening_mood', '20%', 'Итоговая оценка дня'],
    ]
)

h3('12.4 Overall Wellness Score')
code('overall = round((M-Rest + M-Ready + M-Balance) / 3.0)')
body('Метки: Excellent (>= 80), Good (>= 60), Fair (>= 40), Poor (< 40).')

# ── 13 ──────────────────────────────────────────────────────────
h2('13. Почему XGBoost и CatBoost лучше альтернатив')

h3('13.1 Сравнение алгоритмов для нашей задачи')
add_table(
    ['Алгоритм', 'Для табличных данных', 'Плюсы', 'Минусы / почему не выбрали'],
    [
        ['XGBoost', 'Лучший выбор', 'L1/L2 регуляризация, early stopping, высокий R²', 'Требует нормализации и энкодинга'],
        ['CatBoost', 'Лучший для рекомендаций', 'Нативные категории, no target leakage', 'Медленнее обучается'],
        ['Random Forest', 'Хуже', 'Устойчивость к overfitting', 'R² ниже, медленнее на инференсе'],
        ['Linear Regression', 'Не подходит', 'Простота интерпретации', 'Не улавливает нелинейные зависимости'],
        ['Neural Network (MLP)', 'Не подходит', 'Теоретически мощнее', 'Нужно в 10–100x больше данных, не интерпретируем'],
        ['LightGBM', 'Близко к XGBoost', 'Быстрее обучается', 'Хуже с малыми датасетами'],
    ]
)

h3('13.2 Почему не нейронные сети (детально)')
body('Четыре ключевых аргумента против Deep Learning для данной задачи:')
bullet(
    'Размер данных. У нас 12 000 строк (health metrics) и синтетический датасет (рекомендации). '
    'Нейросеть на таком объёме переобучится без массивной аугментации данных.'
)
bullet(
    'Интерпретируемость. XGBoost + SHAP даёт объяснение на уровне признаков '
    '(какой фактор сколько повлиял). Это критично для медицинского/wellness-домена, '
    'где пользователю важно понять "почему".'
)
bullet(
    'Скорость инференса. XGBoost: 2–5 мс, MLP: 50–200 мс при той же точности. '
    'При асинхронных Kafka-запросах это принципиально.'
)
bullet(
    'Научный бенчмарк. Shwartz-Ziv & Armon (2022) показали, что gradient boosting '
    'превосходит Deep Learning на табличных данных в 90% случаев, особенно '
    'при числе примеров < 10 000.'
)

h3('13.3 Преимущество CatBoost перед XGBoost для рекомендаций')
body(
    'Ключевая причина — поле action_type (SLEEP_INCREASE, STRESS_DECREASE и т.д.) является '
    'категориальным признаком. XGBoost требует ручного one-hot encoding, '
    'что создаёт target leakage при стандартном подходе. '
    'CatBoost применяет Ordered Target Statistics (Prokhorenkova et al., 2018) — '
    'это вычисляет статистики таргета без включения текущего примера в расчёт, '
    'исключая утечку информации из обучения в тест. '
    'Результат: R² = 0.966 для предсказания improvement, '
    'что существенно лучше, чем XGBoost с тем же датасетом.'
)

# ── 14 ──────────────────────────────────────────────────────────
h2('14. Интеграция ML-сервиса в общую систему')

h3('14.1 API эндпойнты ML Service')
body('Base URL внутри Docker: http://ml-service:5001 (снаружи: http://localhost:5001)')
add_table(
    ['Метод', 'Путь', 'Описание'],
    [
        ['GET', '/health', 'Статус: { status: "healthy", version: "5.1.0" }'],
        ['POST', '/predict', 'XGBoost: когнитивный скор + состояние + CFI'],
        ['POST', '/recommend/top3', 'CatBoost: топ-3 рекомендации с predicted_improvement'],
        ['POST', '/recommend/single', 'Одна лучшая рекомендация'],
        ['POST', '/health-metrics/calculate', 'Формульный расчёт M-Rest/M-Ready/M-Balance'],
        ['POST', '/health-metrics/ml-predict', 'XGBoost-предсказание (fallback на формулу при недоступности)'],
        ['POST', '/health-metrics/trend', 'Тренд за N дней: improving / declining / stable'],
        ['POST', '/internal/cache/update', 'Обновить per-user кэш (вызывает только Java бэкенд)'],
    ]
)

body('Пример запроса к POST /predict:')
code(
    '{\n'
    '  "sleep_duration": 7.0,\n'
    '  "stress_level": 5.0,\n'
    '  "daily_screen_time": 8.0,\n'
    '  "exercise_frequency": 3.0,\n'
    '  "caffeine_intake": 2.0,\n'
    '  "reaction_time": 300.0,\n'
    '  "memory_test_score": 75.0,\n'
    '  "age": 25.0,\n'
    '  "gender": "Male",\n'
    '  "diet_type": "Non-Vegetarian"\n'
    '}'
)
body('Ответ: cognitive_score (0–100), cognitive_state (High/Medium/Low — порог >75 / 50–75 / <50), cfi_score.')

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ IV
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ IV — ИНФРАСТРУКТУРА')

# ── 15 ──────────────────────────────────────────────────────────
h2('15. Kafka — Асинхронные события')

body(
    'Apache Kafka обеспечивает полное расцепление сервисов. '
    'Весь обмен между CheckIn Service и ML/Health-подсистемой идёт через события, '
    'без синхронных HTTP-вызовов в критическом пути. '
    'Это даёт три ключевых преимущества: ответ пользователю за < 30 мс '
    '(только запись в БД), ML inference (до 200 мс) — асинхронно; '
    'при падении ML сервиса чекин сохранён и метрики пересчитаются позже; '
    'сообщения кейсируются по userId — все события одного пользователя '
    'идут в одну партицию, гарантируя порядок обработки (partition-based ordering).'
)

add_table(
    ['Топик', 'Партиции', 'Продюсер', 'Потребители', 'Когда публикуется'],
    [
        ['checkin.created', '3', 'NBCheckinService', 'HealthMetricsSaver, MLRecommendationConsumer', 'POST /checkins'],
        ['sleep.logged', '3', 'NBCheckinService', 'SleepLogKafkaConsumer', 'POST /sleep'],
        ['game.completed', '3', 'NBCheckinService', 'HealthMetricsKafkaConsumer, ML refresh', 'POST /brain-games/submit'],
        ['character.leveled-up', '3', 'NBCheckinService', 'CharacterProgressionConsumer', 'После level-up'],
    ]
)

h3('15.1 Детальный флоу после создания чекина')
code(
    'POST /checkins  →  HTTP 201  (< 30 мс, ответ пользователю)\n'
    '     │\n'
    '     ▼\n'
    'DailyCheckInService.createCheckIn()  [@Transactional]\n'
    '     ├── Сохраняет DailyCheckIn в БД\n'
    '     ├── Обновляет UserStreak (increment или reset)\n'
    '     ├── Автовыполняет задачу COMPLETE_CHECKIN (+50 XP)\n'
    '     └── Публикует CheckInCreatedApplicationEvent (Spring Event)\n'
    '              │\n'
    '              ▼ (после коммита транзакции, TransactionalKafkaPublisher)\n'
    '     kafkaTemplate.send("checkin.created", userId.toString(), event)\n'
    '              │\n'
    '              ├── HealthMetricsSaver  [groupId: health-metrics-group]\n'
    '              │       └── POST /predict (ML Service)\n'
    '              │               → сохраняет HealthMetrics в БД\n'
    '              │\n'
    '              └── MLRecommendationConsumer  [groupId: ml-recommendations-group]\n'
    '                      └── MLRecommendationCacheService.asyncRefresh()\n'
    '                              └── POST /recommend/top3 (internal=true)\n'
    '                                      → сохраняет в daily_ml_recommendation\n'
    '                                      → обновляет Python кэш: /internal/cache/update'
)

h3('15.2 Graceful Degradation при недоступности ML')
body(
    'Если ML Service недоступен при обработке события checkin.created: '
    'HealthMetricsSaver вычисляет M-Rest/M-Ready/M-Balance по детерминированным формулам '
    '(раздел 12), сохраняет в health_metrics с mlScore = null, '
    'пользователь видит метрики без когнитивного скора. '
    'Когда ML восстанавливается — фронтенд может вызвать '
    'POST /health-metrics/recalculate для принудительного пересчёта.'
)
body(
    'Реализация: kafkaTemplate настроен с StringSerializer для ключей и JsonSerializer '
    'для значений. Ключ = userId (String), что обеспечивает routing в одну партицию '
    'для всех событий одного пользователя.'
)

# ── 16 ──────────────────────────────────────────────────────────
h2('16. Базы данных')

h3('16.1 auth_db (порт 5432)')
add_table(
    ['Таблица', 'Ключевые поля'],
    [
        ['users', 'id, name, email, username, phone, password_hash (BCrypt), is_verified, is_onboarded, refresh_token'],
        ['user_onboarding', 'user_id, sex, height_cm, weight_kg, birth_date, character_id, data_consent'],
        ['verification_tokens', 'token (UUID), user_id, expires_at'],
    ]
)

h3('16.2 checkin_db (порт 5433)')
add_table(
    ['Таблица', 'Назначение'],
    [
        ['daily_check_ins', 'Ежедневные чекины (user_id + date — unique constraint)'],
        ['mood_logs', 'Логи настроения (несколько в день)'],
        ['sleep_logs', 'Записи о сне (bedtime, wake_time, duration_hours, quality_score, deep_sleep_minutes)'],
        ['daily_tasks', 'Ежедневные задачи (task_type, is_completed, completed_at, xp_reward)'],
        ['health_metrics', 'M-Rest, M-Ready, M-Balance, overall, ml_score по датам'],
        ['daily_ml_recommendation', 'Кэш ML-рекомендаций: cognitive_score, recommendations_json, trigger_source'],
        ['user_streaks', 'current_streak, longest_streak, total_checkins, total_xp_earned, last_checkin_date'],
        ['user_rewards', 'Бейджи: reward_type, unlocked_at, xp_bonus, xp_multiplier'],
        ['user_characters', 'character_type, level, xp, happiness, energy'],
        ['brain_game_results', 'Результаты когнитивных игр (game_type, score, difficulty_level, correct_answers)'],
        ['game_sessions', 'Fun-игры: DONUT_GAME, CHARACTER_CARE'],
        ['new_game_sessions', 'Новые игры: DONUT_GAME, NUMBER_SEQUENCE_GAME + attempts_count'],
        ['user_game_stats', 'Агрегированная статистика: total_games, best_score, total_xp'],
    ]
)

h3('16.3 noteai_db (порт 5434)')
add_table(
    ['Таблица', 'Назначение'],
    [
        ['notes', 'Старые заметки (title, content)'],
        ['note_users', 'Пользователи — зеркало из auth_db'],
        ['journal_entries', 'Записи расширенного дневника (title, content, mood_score, tags, is_favorite)'],
    ]
)

# ── 17 ──────────────────────────────────────────────────────────
h2('17. JWT и безопасность')
add_table(
    ['Аспект', 'Решение', 'Стандарт/Основание'],
    [
        ['JWT алгоритм', 'HS256 (HMAC-SHA-256)', 'RFC 7519 — симметричный, секрет в JWT_SECRET'],
        ['Срок токена', '86 400 сек (24 часа)', 'Конфигурируется через JWT_EXPIRATION'],
        ['Refresh токен', 'UUID в таблице users', 'Ротация при каждом /auth/refresh'],
        ['Хэширование паролей', 'BCrypt, cost factor 10', 'NIST SP 800-63B'],
        ['Хранение JWT на iOS', 'Keychain kSecClassGenericPassword', 'OWASP MASVS-STORAGE-1'],
        ['Шифрование Keychain', 'Secure Enclave устройства', 'Недоступно другим приложениям'],
        ['Межсервисная авторизация', 'Shared JWT_SECRET (env var)', 'Каждый сервис валидирует самостоятельно'],
        ['ML Service доступность', 'Только внутри Docker-сети', 'Нет внешнего доступа к ML inference'],
    ]
)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ V — ИГРОВАЯ СИСТЕМА
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ V — ИГРОВАЯ СИСТЕМА')

h2('18. XP, персонажи, стрики, бейджи')

h3('18.1 Начисление XP')
add_table(
    ['Действие', 'XP', 'Частота'],
    [
        ['Ежедневный чекин (COMPLETE_CHECKIN)', '+50', 'Раз в день'],
        ['Написать заметку/дневник (WRITE_NOTE)', '+45', 'Раз в день'],
        ['Когнитивная игра (PLAY_GAME)', '+30–80 (зависит от результата)', 'До 3 раз в день'],
        ['Записать сон (LOG_SLEEP)', '+25', 'Раз в день'],
        ['Записать настроение (LOG_MOOD)', '+20', 'Раз в день'],
        ['Milestone стрик 7 дней', '+200 (единоразово)', 'Один раз'],
        ['Milestone стрик 30 дней', '+500 (единоразово)', 'Один раз'],
    ]
)

h3('18.2 XP-множители')
add_table(
    ['Условие', 'Множитель'],
    [
        ['Стрик >= 7 дней', '× 1.1 (+10% к XP)'],
        ['Стрик >= 30 дней', '× 1.25 (+25% к XP)'],
    ]
)

h3('18.3 Уровни персонажа — экспоненциальная кривая')
body('Формула порогового XP для уровня N:')
code('XP_threshold(N) = N² × 100')
body(
    'Уровень 2 требует 400 XP, уровень 3 — 900, уровень 4 — 1600, уровень 5 — 2500. '
    'Экспоненциальный рост предотвращает инфляцию уровней в начале '
    'и поддерживает долгосрочную вовлечённость (retention).'
)

h3('18.4 Адаптивная сложность игр')
body(
    'Уровень сложности когнитивных игр определяется текущим когнитивным скором из XGBoost. '
    'Это реализация принципа Flow Theory (Csikszentmihalyi, 1990) — '
    '"оптимальный опыт" когда сложность соответствует уровню навыка:'
)
bullet('Скор > 75 — HARD (пользователь в хорошей когнитивной форме)')
bullet('Скор 50–75 — MEDIUM (средняя форма)')
bullet('Скор < 50 — EASY (низкий скор = лёгкие задания для восстановления уверенности)')

h3('18.5 Бейджи (Rewards)')
add_table(
    ['Бейдж', 'Условие разблокировки', 'XP-бонус', 'XP-множитель'],
    [
        ['FIRST_CHECKIN', 'Первый чекин в жизни', '—', '—'],
        ['STREAK_7', '7 последовательных дней с чекином', '+100', '× 1.1'],
        ['STREAK_30', '30 последовательных дней с чекином', '+500', '× 1.25'],
        ['GAME_MASTER', 'Определённое количество игр', '—', '—'],
    ]
)

h3('18.6 Логика стрика')
body(
    'Стрик — количество последовательных дней с выполненным чекином. '
    'Пропуск одного дня сбрасывает стрик в 0. '
    'Логика в DailyCheckInService при создании чекина:'
)
code(
    'if last_checkin_date == today - 1 (вчера, Asia/Almaty):\n'
    '    current_streak += 1\n'
    'else:\n'
    '    current_streak = 1\n'
    '\n'
    'longest_streak = max(longest_streak, current_streak)\n'
    'last_checkin_date = today'
)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ VI — ФЛОУ
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ VI — ПОЛНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ')

h2('19. Примеры флоу от начала до конца')

h3('Флоу 1: Первый запуск приложения')
body('Полный путь нового пользователя от установки до готовности:')
code(
    '1. POST /api/v1/auth/register\n'
    '   → Создаётся пользователь, isOnboarded=false\n'
    '\n'
    '2. POST /api/v1/auth/login\n'
    '   → accessToken + refreshToken → iOS Keychain\n'
    '\n'
    '3. POST /api/v1/onboarding\n'
    '   → Пол, рост, вес, дата рождения, ID персонажа\n'
    '\n'
    '4. POST /api/v1/character/select  (CheckIn Service)\n'
    '   → Выбираем: CAT / DOG / RABBIT / PANDA / FOX'
)

h3('Флоу 2: Ежедневный чекин')
code(
    '1. GET /checkins/today/exists\n'
    '   → { "exists": false, "canCheckIn": true }\n'
    '\n'
    '2. POST /checkins\n'
    '   → { morningMood:4, sleepQuality:7, energyLevel:6, stressLevel:4, ... }\n'
    '   → HTTP 201 за < 30 мс\n'
    '\n'
    '   (Фоново через Kafka, ~ 200 мс):\n'
    '   → Health Metrics пересчитаны: M-Rest=78, M-Ready=72, M-Balance=65\n'
    '   → ML-рекомендации обновлены\n'
    '   → Задача COMPLETE_CHECKIN = выполнена (+50 XP)\n'
    '\n'
    '3. GET /health-metrics/today\n'
    '   → { mRest:78.5, mReady:72.3, mBalance:65.1, overall:71.9 }\n'
    '\n'
    '4. POST /ml/recommendations\n'
    '   → Топ-3 персонализированных рекомендации с predicted_improvement'
)

h3('Флоу 3: Ведение дневника с AI-анализом')
code(
    '1. POST /api/v1/journal\n'
    '   → { title:"Мои мысли", content:"...", moodScore:4 }\n'
    '   → Задача WRITE_NOTE = выполнена, +45 XP\n'
    '\n'
    '2. Пользователь продолжает писать →\n'
    '   PATCH /api/v1/journal/{id}  (автосейв debounce 2–3 сек)\n'
    '   → { content:"...обновлённый текст..." }\n'
    '\n'
    '3. POST /api/v1/journal/{id}/ai/analyze\n'
    '   → { tone, themes, wellnessInsight, suggestion }\n'
    '\n'
    '4. POST /api/v1/journal/ai/chat\n'
    '   → { message:"Как улучшить сон?", noteId:15 }\n'
    '   → AI отвечает с контекстом записи (Groq Llama 3.3-70b)'
)

h3('Флоу 4: Обновление access токена')
code(
    '1. Получаем 401 Unauthorized (токен истёк)\n'
    '\n'
    '2. POST /api/v1/auth/refresh\n'
    '   Body (raw string): <refresh_token>\n'
    '\n'
    '3. Получаем новые accessToken + refreshToken\n'
    '\n'
    '4. Повторяем оригинальный запрос с новым токеном'
)

h3('Флоу 5: Игра → рост персонажа → бейдж')
code(
    '1. POST /brain-games/submit\n'
    '   → { gameType:"NUMBER_SEQUENCE", score:850, durationSeconds:45 }\n'
    '   → Kafka: game.completed → ML-рекомендации обновляются\n'
    '\n'
    '2. GET /brain-games/progression\n'
    '   → { weeklyGames:3, requiredGames:5, xpThisWeek:240 }\n'
    '\n'
    '3. POST /brain-games/level-up  (после набора XP за неделю)\n'
    '   → { level:4, justLeveledUp:true }\n'
    '\n'
    '4. POST /rewards/check\n'
    '   → Проверка и выдача новых бейджей'
)

pagebreak()

# ════════════════════════════════════════════════════════════════
# ЧАСТЬ VII — ДИПЛОМ
# ════════════════════════════════════════════════════════════════
h1('ЧАСТЬ VII — ДЛЯ ЗАЩИТЫ ДИПЛОМА')

h2('20. Что рассказать на защите: бэкенд-часть')

h3('20.1 Ключевые архитектурные решения (аргументы для комиссии)')
h4('Зачем микросервисы, а не монолит?')
bullet('Независимое развёртывание: ML Service обновляется без остановки Auth и CheckIn.')
bullet('Изоляция отказов: падение NoteAI не влияет на чекины и метрики здоровья.')
bullet('Масштабирование: ML Service можно вынести на GPU-сервер отдельно при росте нагрузки.')
bullet('Database-per-Service (Newman, 2015): нет shared schema coupling, независимые миграции.')

h4('Зачем Kafka, а не прямые синхронные HTTP-вызовы?')
bullet('Bounded latency: ответ пользователю всегда < 30 мс (только запись в БД).')
bullet('ML inference до 200 мс — асинхронно, не блокирует пользователя.')
bullet('Resilience: при падении ML пользователь всё равно получает HTTP 201 и видит метрики по формулам.')
bullet('Ordering: все события одного userId идут в одну партицию — исключает race condition.')

h4('Зачем три отдельные БД?')
bullet('Безопасность: auth_db с BCrypt-хэшами изолирован от остальных данных.')
bullet('Независимые миграции: Flyway V1__, V2__... для каждой БД без конфликтов.')
bullet('HikariCP на каждом сервисе: max pool = 10, оптимально для микросервисной нагрузки.')

h3('20.2 Быстрые ответы на вопросы комиссии')
add_table(
    ['Вопрос комиссии', 'Краткий ответ'],
    [
        ['Как хранятся пароли?', 'BCrypt, cost factor 10. Адаптивный — при росте железа время хэширования растёт автоматически.'],
        ['Как работает refresh token?', 'UUID в таблице users. /auth/refresh выдаёт новую пару токенов и обновляет запись в БД.'],
        ['Что происходит при падении ML?', 'Graceful degradation: HealthMetricsSaver считает формулы, mlScore = null, пользователь видит метрики.'],
        ['Как избегаем двойных чекинов?', 'Unique constraint (user_id, date) в таблице daily_check_ins. Повторный POST — 409 Conflict.'],
        ['Как работает стрик?', 'last_checkin_date сравнивается с today-1 по Asia/Almaty. Если совпадает — increment, иначе reset=1.'],
        ['Как Kafka гарантирует порядок?', 'Ключ сообщения = userId String. Kafka роутит все сообщения с одним ключом в одну партицию.'],
        ['Как автовыполняются задачи?', 'Spring ApplicationEvent (CheckInCreatedEvent) → @EventListener в DailyTaskService.completeTask().'],
        ['Сколько контейнеров?', '10: 3 Spring Boot + 3 PostgreSQL + ML Flask + Kafka + Zookeeper + pgAdmin.'],
        ['Какой LLM используется в NoteAI?', 'Groq API, модель llama-3.3-70b-versatile, temperature=0.3 для детерминированности.'],
        ['Как тестируются API?', 'Swagger UI на /swagger-ui.html каждого сервиса + ручная верификация через /auth/verify-user-manual.'],
    ]
)

# ── 21 ──────────────────────────────────────────────────────────
h2('21. Что рассказать на защите: ML-часть')

h3('21.1 Структура рассказа о ML (8–10 минут)')

h4('Шаг 1 — Постановка задачи')
body(
    'Нам нужно предсказывать когнитивное состояние человека из данных ежедневного '
    'самоотчёта. Это объединяет две классические ML-задачи: регрессию '
    '(скор 0–100 — непрерывная величина) и многоклассовую классификацию '
    '(три категории: High, Medium, Low).'
)

h4('Шаг 2 — Данные')
body(
    'Исходный датасет: human_cognitive_performance.csv с физиологически валидированными '
    'данными. 10 исходных признаков плюс 7 производных = 17 итоговых признаков (FINAL_FEATURES). '
    'Для Health Metrics — 12 000 синтетических записей с target values по физиологическим формулам. '
    'Для рекомендаций — 50 000 записей с рассчитанным improvement для каждого сценария.'
)

h4('Шаг 3 — Feature Engineering')
body(
    'Ключевой производный признак — CFI (Cognitive Fatigue Index): взвешенная комбинация '
    'реакции (30%), стресса (25%), дефицита сна (25%) и экранного времени (20%). '
    'Это composite feature улавливает взаимодействия, которые сырые признаки '
    'не могут выразить по отдельности. Основание: Killgore (2010) — '
    'когнитивное утомление является мультифакторным феноменом.'
)

h4('Шаг 4 — Выбор алгоритма')
body(
    'XGBoost выбран как бенчмарк-победитель для табличных данных '
    '(Chen & Guestrin 2016 — 29 из 32 побед на Kaggle). '
    'CatBoost — для рекомендаций: нативно обрабатывает категориальный признак action_type '
    'без target leakage через Ordered Target Statistics (Prokhorenkova et al., 2018).'
)

h4('Шаг 5 — Обучение')
body(
    'Разбивка: 70/10/20 (train/validation/test), стратифицированная по Cognitive_State. '
    'StandardScaler: fit ТОЛЬКО на train, transform на val/test — исключает data leakage. '
    'Early stopping: 30 раундов без улучшения validation RMSE. '
    'L1 (reg_alpha=0.1) и L2 (reg_lambda=1.0) регуляризация.'
)

h4('Шаг 6 — Результаты')
body(
    'XGBoost Regressor: R²=0.9990, RMSE=0.73 — ошибка менее 1 балла из 100. '
    'XGBoost Classifier: Accuracy=97.9%. '
    'CatBoost рекомендации: R²=0.966, RMSE=1.81 — ошибка предсказания improvement менее 2 баллов. '
    'Health Metrics: M-Ready R²=0.971, MAE=2.02.'
)

h4('Шаг 7 — Интеграция')
body(
    'Модели загружаются в память при старте Flask — zero disk I/O при инференсе. '
    'XGBoost: 2–5 мс, CatBoost: 10–20 мс. ML Service изолирован внутри Docker-сети. '
    'При недоступности — автоматический fallback на детерминированные формулы (graceful degradation).'
)

h3('21.2 Ответы на сложные вопросы комиссии по ML')
add_table(
    ['Вопрос', 'Полный ответ'],
    [
        ['Почему синтетические данные?',
         'Реальные размеченные данные по когнитивной производительности недоступны публично '
         'в нужном объёме. Синтетика создана на основе физиологически валидированных '
         'распределений из рецензируемой литературы (Walker 2017, McEwen 2007, Killgore 2010). '
         'Это стандартная практика в medical ML для первичных прототипов.'],
        ['Откуда такой высокий R²=0.999?',
         'Модель обучена на синтетических данных с заложенными математическими зависимостями — '
         'поэтому она их улавливает почти идеально. На реальных данных, '
         'с индивидуальным шумом и вариабельностью, R² составит ~0.85–0.92. '
         'Это честно и указано в документации.'],
        ['Как проверяли переобучение?',
         'Три уровня защиты: (1) стратифицированный split 70/10/20, '
         '(2) early stopping по validation set, '
         '(3) изолированный test set, не участвовавший в подборе параметров. '
         'Разница train RMSE / val RMSE незначительна — overfitting нет.'],
        ['Почему не нейронные сети?',
         'Shwartz-Ziv & Armon (2022): gradient boosting превосходит DL на табличных данных '
         'в 90% случаев. У нас 12 000 строк — слишком мало для MLP. '
         'XGBoost + SHAP интерпретируем, что важно для wellness-домена. '
         'Инференс 2–5 мс vs 50–200 мс у нейросети при том же качестве.'],
        ['Что такое CFI?',
         'Cognitive Fatigue Index — авторский составной признак, основанный на '
         'исследованиях Killgore (2010). Четыре компонента: '
         '30% время реакции + 25% стресс + 25% дефицит сна + 20% избыток экрана. '
         'Моделирует кумулятивное когнитивное утомление от нескольких факторов.'],
        ['Как интерпретируете модель?',
         'SHAP (SHapley Additive exPlanations) — стандартный инструмент для XGBoost. '
         'Для нашей модели Sleep_Duration и Stress_Level имеют наибольшую важность, '
         'что соответствует научной литературе.'],
        ['Почему три отдельные модели для метрик?',
         'M-Rest определяется преимущественно сном, M-Balance — стрессом и lifestyle, '
         'M-Ready — энергией и утренним настроением. Одна модель с тремя таргетами '
         'не сможет специализироваться под профиль каждой метрики. '
         'Три отдельные модели дают лучший R² для каждой.'],
    ]
)

# ── 22 ──────────────────────────────────────────────────────────
h2('22. Сравнение с аналогами и научное обоснование')

h3('22.1 Конкурентная карта')
add_table(
    ['Приложение', 'ML-скор', 'Персонал. рекомендации', 'AI дневник', 'Геймификация', 'Трекинг сна+настроения'],
    [
        ['NeuroBalance (наш)', 'XGBoost (0.999 R²)', 'CatBoost (0.966 R²)', 'Groq Llama 3.3-70b', 'XP + стрики + бейджи', 'Да (полный)'],
        ['Headspace', 'Нет', 'Контент-based', 'Нет', 'Нет', 'Нет'],
        ['Calm', 'Нет', 'Контент-based', 'Нет', 'Нет', 'Нет'],
        ['Daylio', 'Нет', 'Нет', 'Нет', 'Нет', 'Базовый'],
        ['Apple Health', 'Нет', 'Нет', 'Нет', 'Нет', 'Частично'],
        ['Woebot (AI чат-бот)', 'Нет', 'Rule-based', 'CBT-чат', 'Нет', 'Нет'],
    ]
)
body(
    'Ключевое конкурентное преимущество: NeuroBalance — единственное приложение, '
    'объединяющее ML-driven когнитивный скор + персонализированные рекомендации '
    'с научным обоснованием + AI-дневник + геймификацию в одном продукте.'
)

h3('22.2 Научное обоснование каждого компонента системы')
add_table(
    ['Исследование', 'Что использовано в системе'],
    [
        ['Killgore (2010) — Sleep Deprivation and Cognition',
         'Основа для M-Rest, CFI index, веса sleep_hours в XGBoost'],
        ['Walker (2017) — Why We Sleep',
         'Норма 8 часов, REM/deep sleep компоненты, scientific_basis для SLEEP_INCREASE'],
        ['McEwen (2007) — Allostatic Load',
         'Стресс — ключевой компонент M-Balance (40%), scientific_basis для STRESS_DECREASE'],
        ['Ratey (2008) — Spark: BDNF and Exercise',
         'Физическая активность → BDNF (нейропластичность), scientific_basis для EXERCISE_INCREASE'],
        ['Newport (2016) — Digital Minimalism',
         'Экранное время > 8 часов = cognitive fragmentation, scientific_basis для SCREEN_DECREASE'],
        ['Stults-Kolehmainen & Sinha (2014)',
         'Упражнения + социальная активность + питание = формула M-Balance'],
        ['Csikszentmihalyi (1990) — Flow Theory',
         'Адаптивная сложность игр по когнитивному скору (Easy/Medium/Hard)'],
        ['Chen & Guestrin (2016) — XGBoost paper',
         'Обоснование выбора XGBoost (29/32 Kaggle wins)'],
        ['Prokhorenkova et al. (2018) — CatBoost paper',
         'Обоснование CatBoost для категориальных признаков (action_type)'],
        ['Shwartz-Ziv & Armon (2022) — Tabular Data Benchmark',
         'Обоснование gradient boosting vs Deep Learning для табличных данных'],
        ['Newman (2015) — Building Microservices',
         'Database-per-Service pattern, обоснование изолированных БД'],
    ]
)

# ── 23 ──────────────────────────────────────────────────────────
h2('23. Структура слайдов для презентации')

add_table(
    ['Слайд', 'Заголовок', 'Ключевой контент'],
    [
        ['1', 'Проблема и рынок',
         'Рынок wellness-приложений $5.2 млрд (2023), CAGR 16.5%. '
         'Разрыв: нет ML-driven когнитивного индекса из ежедневных данных.'],
        ['2', 'Решение: NeuroBalance',
         'Что такое когнитивный скор. 4 уникальные фичи vs конкурентов. '
         'Целевая аудитория. Бизнес-модель (Freemium + B2B).'],
        ['3', 'Микросервисная архитектура',
         'Диаграмма 6 сервисов + Kafka + Docker Compose. '
         '10 контейнеров, 3 БД. Почему не монолит.'],
        ['4', 'Поток данных: чекин → метрики',
         'Диаграмма: POST /checkins → Kafka → ML → HealthMetrics → кэш. '
         'Latency: < 30 мс sync + 200 мс async.'],
        ['5', 'Безопасность и аутентификация',
         'JWT HS256, BCrypt cost=10, Keychain iOS. '
         'JwtAuthenticationFilter flow.'],
        ['6', 'ML: постановка задачи и данные',
         '6 моделей, 3 задачи. Датасеты. 10 исходных + 7 engineered = 17 признаков. '
         'CFI формула.'],
        ['7', 'XGBoost: когнитивный скор',
         'Гиперпараметры. Split 70/10/20. Метрики: R²=0.9990, RMSE=0.73. '
         'SHAP — важность признаков.'],
        ['8', 'CatBoost: рекомендации',
         '4 типа рекомендаций. Ordered Target Statistics. '
         'Метрики: R²=0.966, RMSE=1.81. SmartRecommendationSelector.'],
        ['9', 'Health Metrics: формулы и модели',
         'M-Rest, M-Ready, M-Balance — формулы с весами и обоснованием. '
         'R² каждой модели (0.917 / 0.971 / 0.967). Метки.'],
        ['10', 'Сравнение с конкурентами',
         'Таблица: NeuroBalance vs Headspace vs Calm vs Daylio vs Apple Health.'],
        ['11', 'Демонстрация приложения',
         'Скриншоты iOS: онбординг, чекин, дашборд с 3 кольцами, рекомендации, дневник с AI.'],
        ['12', 'Итоги и планы развития',
         'Достигнутые метрики. B2B tier. Wearables интеграция. Корпоративные дашборды.'],
    ]
)

body('Ключевые числа для слайдов — всегда держать наготове:')
code(
    'XGBoost Regressor:      R² = 0.999,  RMSE = 0.73  (из 100 баллов)\n'
    'XGBoost Classifier:     Accuracy = 97.9%\n'
    'CatBoost рекомендации:  R² = 0.966,  RMSE = 1.81\n'
    'M-Rest:    R² = 0.917,  MAE = 2.01\n'
    'M-Ready:   R² = 0.971,  MAE = 2.02\n'
    'M-Balance: R² = 0.967,  MAE = 2.13\n'
    '\n'
    'API latency:   < 30 мс (чекин sync) + ~200 мс (ML async)\n'
    'ML inference:  2–5 мс (XGBoost),  10–20 мс (CatBoost)\n'
    'Признаков:     17 финальных (10 исходных + 7 engineered)\n'
    'Датасет HM:    12 000 записей\n'
    'Датасет Rec:   ~50 000 записей\n'
    'Docker:        10 контейнеров\n'
    'Сервисов:      6 (3 Java + ML Flask + Kafka + Zookeeper)\n'
    'LLM:           Groq Llama 3.3-70b-versatile, temperature=0.3'
)

# ── Финальный абзац ─────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Документ обновлён: 2026-06-07  ·  Версия 3.0  ·  NeuroBalance Backend')
run.font.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Сохранение ──────────────────────────────────────────────────
output = '/Users/amangeldimadina/IdeaProjects/NeuroBalanceBackend/NeuroBalance_Technical_Documentation.docx'
doc.save(output)
print(f'Saved: {output}')
